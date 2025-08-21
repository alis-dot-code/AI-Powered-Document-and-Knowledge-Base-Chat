"""DOCX parser using python-docx."""
from dataclasses import dataclass, field


@dataclass
class ParsedDocument:
    pages: list  # list of dicts {page_number, text} — docx has no real pages
    page_count: int = 1

    @property
    def full_text(self) -> str:
        return "\n\n".join(p["text"] for p in self.pages if p["text"].strip())


def parse(file_bytes: bytes) -> ParsedDocument:
    import io
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)

    # Tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        text += "\n\n" + "\n".join(rows)

    return ParsedDocument(pages=[{"page_number": 1, "text": text.strip()}])
